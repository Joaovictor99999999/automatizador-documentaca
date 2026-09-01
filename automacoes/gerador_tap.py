from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx2pdf import convert
import os
from google import genai
from google.genai import types
from dotenv import load_dotenv

# 1. A FUNÇÃO AGORA RECEBE OS DADOS AO INVÉS DE PERGUNTAR
def rodar_automacao_tap(dados_basicos, assinaturas, contexto, atividades_selecionadas, participantes_selecionados):

    print("--- 🛠️ INICIALIZANDO MOTOR OLS (TAP) ---")
    NOME_TEMPLATE = "templates/cabeçalho_OLS.docx" # <-- AJUSTE AQUI SE SUA PASTA FOR OUTRA

    # 2. BANCO DE DADOS DA MATRIZ (Consulta rápida)
    matriz_db = {
        "Materiais e Equipamentos": {"setor": "Projetos", "sub": "Compras", "resp": "Fabio/Adriano"},
        "Logistica": {"setor": "Projetos", "sub": "Compras", "resp": "Fabio/Adriano"},
        "Montagem de rack e Rede BR": {"setor": "Operações", "sub": "TI", "resp": "André/Alberto"},
        "Sistema Satelitais – LEO": {"setor": "Operações", "sub": "Hub", "resp": "André/Jairo"},
        "Sistema Satelitais – GEO": {"setor": "Operações", "sub": "Hub", "resp": "André/Jairo"},
        "Sistema de CFTV": {"setor": "Operações", "sub": "TI", "resp": "André/Alberto"},
        "Sistema de replicação de Imagens": {"setor": "Operações", "sub": "", "resp": "André"},
        "Sistema de telefonia": {"setor": "Operações", "sub": "TI", "resp": "André/Alberto"},
        "Sistema de TV UHF": {"setor": "Operações", "sub": "", "resp": "André"},
        "Sistema de LTE": {"setor": "Operações", "sub": "", "resp": "André"},
    }

    # 3. PREPARA O TEXTO PARA A IA E PRO WORD
    texto_atividades = ", ".join(atividades_selecionadas)
    
    # Transforma "Gestor do Projeto (@carlosdamatta)" em "Gestor do Projeto: @carlosdamatta;"
    linhas_formatadas = []
    for part in participantes_selecionados:
        cargo, nome = part.replace(")", "").split(" (")
        linhas_formatadas.append(f"{cargo}: {nome};")
    texto_equipe_bloco = "\n".join(linhas_formatadas)

    # =========================================================
    # 4. INTELIGÊNCIA ARTIFICIAL 
    # =========================================================
    print("🧠 Solicitando redação à IA...")
    prompt = f"""
    Você é um especialista na elaboração de Termos de Abertura de Projeto (TAP) da empresa OLS.
    O projeto NÃO é necessariamente naval. Adapte automaticamente o vocabulário ao tipo de empreendimento.
    
    Dados:
    Tipo: {contexto['tipo']}
    Nome: {contexto['nome_emp']}
    Objetivo: {contexto['objetivo']}
    Descrição: {contexto['descricao']}
    Atividades: {texto_atividades}

    Escreva APENAS as três seções abaixo. Sem explicações, sem listas.
    O escopo deve terminar com: "Não contempla nesse projeto a confecção de infraestrutura por parte da OLS."
    
    [OBJETIVO]
    Até 30 palavras.
    [FINALIDADE]
    Até duas frases.
    [ESCOPO]
    Até duas frases curtas.
    """

    try:
        load_dotenv()
        
        api_key = os.getenv("GOOGLE_API_KEY")
        
        # 1. Inicializa o cliente com a sua chave
        client = genai.Client(api_key="api_key")
        # 2. Configura a criatividade e tamanho da resposta
        configuracoes = types.GenerateContentConfig(
            temperature=0.3,
        )

        # 3. Chama o modelo
        resposta = client.models.generate_content(
            model='gemini-3.6-flash',
            contents=prompt,
            config=configuracoes
        )
        
        resultado_ia = resposta.text.strip()
        
        print("\n--- 📝 RESPOSTA BRUTA DA IA ---")
        print(resultado_ia)
        print("-------------------------------\n")
        
    except Exception as e:
        print(f"⚠️ Erro ao chamar IA: {e}")
        resultado_ia = ""

    # Extração das tags
    def extrair_secao(texto, tag_atual, tags_proximas=None):
        if tag_atual not in texto: return ""
        parte = texto.split(tag_atual, 1)[1]
        if tags_proximas:
            pontos_corte = [parte.index(t) for t in tags_proximas if t in parte]
            if pontos_corte: parte = parte[:min(pontos_corte)]
        return parte.strip(" :\n")

    objetivo_texto = extrair_secao(resultado_ia, "[OBJETIVO]", ["[FINALIDADE]", "[FINALID]"])
    finalidade_texto = extrair_secao(resultado_ia, "[FINALIDADE]", ["[ESCOPO]", "[ESCOP]", "[ESCO]"])
    escopo_texto = extrair_secao(resultado_ia, "[ESCOPO]") or extrair_secao(resultado_ia, "[ESCOP]")

    # Fallback de Segurança
    if not objetivo_texto: objetivo_texto = f"{contexto['objetivo']} no empreendimento {contexto['nome_emp']}."
    if not finalidade_texto: finalidade_texto = "Projeto desenvolvido para atender adequações técnicas."
    if not escopo_texto: escopo_texto = f"Atualização de sistemas. Não contempla nesse projeto a confecção de infraestrutura por parte da OLS."

    equipe_recursos_texto = (
        "Para a realização desse projeto, será necessário a participação do Team de TI OLS, "
        "equipe de HUB OLS, equipe de Operações OLS, sendo necessário a utilização dos recursos "
        "como: 2 técnicos OLS abordo, apoio remoto ao menos 1 Analista de TI, 1 Analista de Hub, "
        "além de apoio logístico, também o team de Suprimentos, uso de ferramentas e insumos e "
        "verba para as despesas envolvidas para essa atividade."
    )

    # =========================================================
    # 5. CONSTRUÇÃO DO WORD
    # =========================================================
    print("💾 Montando o documento Word...")
    doc = Document(NOME_TEMPLATE) if os.path.exists(NOME_TEMPLATE) else Document()
    
    if doc.tables:
        tabela = doc.tables[0]
        tabela.cell(0, 5).text = f"Cód: {dados_basicos['codigo']}"
        tabela.cell(2, 0).text = f"Nome do Projeto: {dados_basicos['projeto']}"
        tabela.cell(2, 5).text = f"Versão: {dados_basicos['versao']}"
        tabela.cell(3, 0).text = f"Número do Projeto: {dados_basicos['ccv']}"
        tabela.cell(3, 3).text = f"Data de abertura: {dados_basicos['data_abertura']}"
        
        tabela.cell(4, 0).text = f"Elaboração: {assinaturas['elab_data']}\nElaborado por: {assinaturas['elab_nome']}\nCargo: {assinaturas['elab_cargo']}"
        tabela.cell(4, 2).text = f"Revisão: {assinaturas['rev_data']}\nRevisado por: {assinaturas['rev_nome']}\nCargo: {assinaturas['rev_cargo']}"
        tabela.cell(4, 4).text = f"Aprovação: {assinaturas['aprov_data']}\nAprovado por: {assinaturas['aprov_nome']}\nCargo: {assinaturas['aprov_cargo']}"

    section = doc.sections[0]
    section.top_margin, section.bottom_margin = Cm(3.5), Cm(3)
    section.left_margin, section.right_margin = Cm(2), Cm(2)
    doc.styles['Normal'].font.name, doc.styles['Normal'].font.size = 'Calibri', Pt(10)

    # 🛠️ CORREÇÃO 1: Adicionado espaçamento antes do título e depois do texto
    def add_secao(titulo, texto, justificar=True):
        p_titulo = doc.add_paragraph()
        p_titulo.paragraph_format.space_before = Pt(18) # Dá um respiro de 18 pontos antes
        p_titulo.add_run(titulo).bold = True
        
        p_texto = doc.add_paragraph(texto)
        if justificar: p_texto.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_texto.paragraph_format.space_after = Pt(12) # Dá um respiro de 6 pontos depois

    add_secao("1- Objetivos", objetivo_texto)
    add_secao("2- Cronograma", f"Início estimado: {dados_basicos['inicio']}\nTérmino estimado: {dados_basicos['fim']}", justificar=False)
    add_secao("3- Finalidade", finalidade_texto)
    add_secao("4- Escopo", escopo_texto)
    add_secao("5- Equipe e recursos", equipe_recursos_texto)

    doc.add_page_break()
    h6 = doc.add_paragraph()
    h6.paragraph_format.space_before = Pt(24)
    h6.add_run("6- Matriz de Responsabilidade, Participantes e Aprovadores").bold = True
    
    # 🛠️ CORREÇÃO 2: Criando uma Tabela Real com bordas (Grid) ao invés de usar Tabs
    tabela_matriz = doc.add_table(rows=1, cols=3)
    
    # Força a tabela a ter todas as bordas desenhadas (Lógica de XML segura)
    candidatos_estilo = ['Table Grid', 'TableGrid', 'Tabela com Grade', 'Tabela com grade']
    for nome_estilo in candidatos_estilo:
        try:
            tabela_matriz.style = nome_estilo
            break
        except KeyError:
            continue
    else:
        # Se nenhum estilo padrão funcionar, desenha na força via XML
        tblPr = tabela_matriz._tbl.tblPr
        borders = OxmlElement('w:tblBorders')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '4')
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), '000000')
            borders.append(el)
        tblPr.append(borders)

    # Preenche o cabeçalho da tabela
    hdr_cells = tabela_matriz.rows[0].cells
    hdr_cells[0].text = 'ATIVIDADE'
    hdr_cells[1].text = 'SETOR'
    hdr_cells[2].text = 'RESPONSÁVEL'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True

    # Preenche as linhas com os dados selecionados
    for atv in atividades_selecionadas:
        info = matriz_db[atv]
        setor_completo = f"{info['setor']} / {info['sub']}" if info['sub'] else info['setor']
        
        row_cells = tabela_matriz.add_row().cells
        row_cells[0].text = atv
        row_cells[1].text = setor_completo
        row_cells[2].text = info['resp']

    p_separador = doc.add_paragraph()
    p_separador.paragraph_format.space_before = Pt(30) # Espaço extra após a tabela
    
    for part in participantes_selecionados:
        cargo, nome = part.replace(")", "").split(" (")
        p = doc.add_paragraph()
        p.add_run(f"{cargo}: ").bold = True
        p.add_run(f"{nome};")

    nome_base = f"TAP_{contexto['nome_emp'].replace(' ', '_')}"
    if not contexto['nome_emp']: 
        nome_base = "TAP_Vazio"
        
    doc.save(f"{nome_base}.docx")
    
    try:
        convert(f"{nome_base}.docx", f"{nome_base}.pdf")
    except Exception:
        pass