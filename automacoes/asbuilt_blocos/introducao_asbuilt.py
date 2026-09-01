import time
from google import genai
from google.genai import types
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from automacoes.asbuilt_blocos.config_layout import Layout
import os
from dotenv import load_dotenv

load_dotenv()

api_key = os.getenv("GOOGLE_API_KEY")

def gerar_topicos_com_ia(anotacoes_brutas):
    client = genai.Client(api_key=api_key)
    prompt = f"Transforme as anotações soltas abaixo em uma lista de tópicos...\n{anotacoes_brutas}"
    for tentativa in range(3):
        try:
            response = client.models.generate_content(model='gemini-3.6-flash', contents=prompt)
            return response.text.replace("**", "").strip()
        except Exception as e:
            if "429" in str(e) or "RESOURCE_EXHAUSTED" in str(e):
                time.sleep(20)
            else:
                return "- Erro de conexão com a IA."
    return "- Não foi possível gerar os tópicos."

def adicionar_titulo(doc, texto, forcar_respiro_topo=False):
    if forcar_respiro_topo:
        p_invisivel = doc.add_paragraph()
        p_invisivel.paragraph_format.space_after = Layout.RESPIRO_TOPO_PAGINA
        
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = True
    run.italic = True
    run.font.size = Layout.TAMANHO_TITULO_PRINCIPAL
    
    if not forcar_respiro_topo:
        p.paragraph_format.space_before = Layout.ESPACO_ENTRE_BLOCOS
        
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.left_indent = Layout.RECUO_ESQUERDA_NORMAL
    p.paragraph_format.right_indent = Layout.RECUO_DIREITA_NORMAL

# 🚨 A FUNÇÃO AGORA RECEBE 'dados_introducao' DA INTERFACE
def gerar_bloco_introducao(doc, dados_introducao):
    print("\n==================================================")
    print("🤖 BLOCO 1, 2 E 3 - INTRODUÇÃO E DESCRIÇÃO")
    print("==================================================")
    
    # 1. Pega os dados do pacote da interface (com valores padrão de segurança)
    anotacoes_tecnico = dados_introducao.get("anotacoes") or "Sem anotações fornecidas."
    sistemas_unifilar = dados_introducao.get("sistemas_unifilar") or "rede, voz, Vsat LEO, GEO e CFTV"
    
    # 2. Envia para o Gemini formatar
    texto_descricao = gerar_topicos_com_ia(anotacoes_tecnico)

    # --- ITEM 1: DESCRIÇÃO ---
    adicionar_titulo(doc, "1. Descrição", forcar_respiro_topo=True)
    
    p_intro = doc.add_paragraph("As atividades descritas neste documento têm como finalidade atender à TIC_ET-0600.00-5510-760-PPT-542 para a modernização e instalação dos sistemas de telecomunicações nas embarcações da companhia.")
    p_intro.paragraph_format.space_after = Pt(12)
    p_intro.paragraph_format.left_indent = Layout.RECUO_ESQUERDA_NORMAL
    p_intro.paragraph_format.right_indent = Layout.RECUO_DIREITA_NORMAL 
    
    p_contemplado = doc.add_paragraph("Está contemplado neste projeto as seguintes atividades:")
    p_contemplado.paragraph_format.left_indent = Layout.RECUO_ESQUERDA_NORMAL
    p_contemplado.paragraph_format.right_indent = Layout.RECUO_DIREITA_NORMAL
    
    for topico in texto_descricao.split('\n'):
        topico_limpo = topico.strip().lstrip('-').strip()
        if topico_limpo:
            p_bullet = doc.add_paragraph(f"•  {topico_limpo}")
            p_bullet.paragraph_format.left_indent = Layout.RECUO_ESQUERDA_TOPICO
            p_bullet.paragraph_format.right_indent = Layout.RECUO_DIREITA_NORMAL 

    # --- ITEM 2: PLANTAS E DESENHOS TÉCNICOS ---
    adicionar_titulo(doc, "2. Plantas e desenhos técnicos;")
    
    p_vide = doc.add_paragraph()
    p_vide.add_run("Vide anexos:").bold = True
    p_vide.paragraph_format.left_indent = Layout.RECUO_ESQUERDA_NORMAL
    
    for anexo in [f"Anexo I: Unifilar dos sistemas de {sistemas_unifilar};", 
                  "Anexo II: Diagrama By-face do rack de Telecom;", 
                  "Anexo III: ART atestado de responsabilidade técnica"]:
        p_anexo = doc.add_paragraph(anexo)
        p_anexo.paragraph_format.left_indent = Layout.RECUO_ESQUERDA_NORMAL
        p_anexo.paragraph_format.right_indent = Layout.RECUO_DIREITA_NORMAL

    # --- ITEM 3: ESPECIFICAÇÕES TÉCNICAS ---
    adicionar_titulo(doc, "3. Especificações técnicas")
    p_esp = doc.add_paragraph("O tema aborda os tipos de equipamento para cada sistema bem como a forma de instalação atendendo aos padrões definidos na ET TIC_ET-0600.00-5510-760-PPT-542.")
    p_esp.paragraph_format.left_indent = Layout.RECUO_ESQUERDA_NORMAL
    p_esp.paragraph_format.right_indent = Layout.RECUO_DIREITA_NORMAL

    print("✅ Bloco de Introdução com limites laterais gerado com sucesso!")
    return doc