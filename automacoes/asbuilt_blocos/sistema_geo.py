from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from automacoes.asbuilt_blocos.config_layout import Layout
from automacoes.asbuilt_blocos.motor_fotos_locais import anexar_fotos_do_bloco

def adicionar_bordas_tabela(tabela):
    """Desenha bordas pretas simples contornando e dividindo a tabela via XML."""
    tblPr = tabela._tbl.tblPr
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)
        
    for borda in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{borda}')
        b.set(qn('w:val'), 'single') 
        b.set(qn('w:sz'), '4')       
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000') 
        tblBorders.append(b)

def gerar_bloco_geo(doc, qtd_antenas, lista_fotos):
    print("\n==================================================")
    print("📡 BLOCO 6 - SISTEMA GEO (INTELLIAN / SAILOR)")
    print("==================================================")

    # 1. Quebra de página isolando o capítulo
    doc.add_page_break()
    p_invisivel = doc.add_paragraph()
    p_invisivel.paragraph_format.space_after = Layout.RESPIRO_TOPO_PAGINA

    # 2. Títulos
    p_titulo = doc.add_paragraph()
    run_titulo = p_titulo.add_run("6. Sistema GEO")
    run_titulo.bold = True
    run_titulo.italic = True
    run_titulo.font.size = Layout.TAMANHO_TITULO_PRINCIPAL
    p_titulo.paragraph_format.left_indent = Layout.RECUO_ESQUERDA_NORMAL
    p_titulo.paragraph_format.space_after = Pt(12)

    p_subtitulo = doc.add_paragraph()
    run_sub = p_subtitulo.add_run("Instalação")
    run_sub.bold = True
    p_subtitulo.paragraph_format.left_indent = Layout.RECUO_ESQUERDA_NORMAL
    p_subtitulo.paragraph_format.space_after = Pt(6)

    # 3. Lógica de Matrizes (Tabelas) e Texto
    dados_comuns = [
        ("MODEM", "NEWTEC", "MDM2510", "CBO/BR")
    ]
    
    dados_antena_1 = [
        ("ANTENA_KU_BAND", "INTELLIAN", "vX100", "BR (Petrobras)"),
        ("CONTROLADORA ACU", "INTELLIAN", "vX100", "BR"),
        ("BUC", "xxxx", "xxxx", "xxxx"),
        ("LNB", "xxxx", "xxxx", "xxxx")
    ]
    
    dados_antena_2 = [
        ("ANTENA_KU_BAND", "SAILOR", "VSAT Series", "CBO"),
        ("CONTROLADORA ACU", "SAILOR", "ACU Unit", "CBO"),
        ("BUC", "xxxx", "xxxx", "xxxx"),
        ("LNB", "xxxx", "xxxx", "xxxx")
    ]

    if qtd_antenas == "2":
        texto_instalacao = (
            "Da parte do sistema VSAT GEO, foi realizada a substituição e comissionamento das antenas existentes "
            "seguindo os critérios de redundância e operabilidade do contrato. A antena principal homologada para "
            "a rede Petrobras (BR) passou a ser o modelo Intellian vX100, instalada no bordo correspondente. "
            "Para o link dedicado da embarcação (CBO), foi integrada e comissionada uma antena Sailor (Cobham). "
            "Toda a infraestrutura de cabos coaxiais e conectores de radiofrequência foi inspecionada, testada "
            "e adequada para garantir a atenuação mínima de sinal e integridade dos transceptores."
        )
        dados_tabela_infra = dados_antena_1 + dados_antena_2 + dados_comuns
        
        # Matriz da Tabela 2 (Equipamentos S/N GEO)
        dados_tabela_equip = [
            ("Antena (BR)", "INTELLIAN", "vX100", "xxxxxxx", "xxxxxxx"),
            ("Antena (CBO)", "SAILOR", "VSAT Series", "xxxxxxx", "xxxxxxx")
        ]
    else:
        texto_instalacao = (
            "Da parte do sistema VSAT GEO, foi realizada a substituição e comissionamento da antena existente "
            "seguindo os critérios de operabilidade do contrato. A antena principal homologada para "
            "a rede Petrobras (BR) passou a ser o modelo Intellian vX100, instalada no bordo correspondente. "
            "Toda a infraestrutura de cabos coaxiais e conectores de radiofrequência foi inspecionada, testada "
            "e adequada para garantir a atenuação mínima de sinal e integridade dos transceptores."
        )
        dados_tabela_infra = dados_antena_1 + dados_comuns
        
        # Matriz da Tabela 2 (Equipamentos S/N GEO)
        dados_tabela_equip = [
            ("Antena (BR)", "INTELLIAN", "vX100", "xxxxxxx", "xxxxxxx")
        ]

    # 4. Injetando o Texto
    p_texto = doc.add_paragraph(texto_instalacao)
    p_texto.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_texto.paragraph_format.left_indent = Layout.RECUO_ESQUERDA_NORMAL
    p_texto.paragraph_format.right_indent = Layout.RECUO_DIREITA_NORMAL
    p_texto.paragraph_format.space_after = Pt(18)

    # ====================================================
    # TABELA 1: INFRAESTRUTURA GEO
    # ====================================================
    p_bullet_tabela = doc.add_paragraph()
    p_bullet_tabela.add_run("•  Sistema Satelitais – GEO").bold = True
    p_bullet_tabela.paragraph_format.left_indent = Layout.RECUO_ESQUERDA_TOPICO
    p_bullet_tabela.paragraph_format.space_after = Pt(12)

    tabela1 = doc.add_table(rows=len(dados_tabela_infra) + 1, cols=4)
    tabela1.alignment = WD_TABLE_ALIGNMENT.CENTER
    adicionar_bordas_tabela(tabela1)

    cabecalhos1 = ["EQUIPAMENTO", "FABRICANTE", "MODELO", "APLICAÇÃO"]
    for i, texto in enumerate(cabecalhos1):
        cell = tabela1.rows[0].cells[i]
        cell.text = texto
        cell.paragraphs[0].runs[0].bold = True

    for i, linha_dados in enumerate(dados_tabela_infra, start=1):
        for col, valor in enumerate(linha_dados):
            tabela1.rows[i].cells[col].text = str(valor)

    for row in tabela1.rows:
        for cell in row.cells: cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("") # Respiro longo entre as tabelas

    # ====================================================
    # TABELA 2: IDENTIFICAÇÃO DOS EQUIPAMENTOS (S/N e IMEI)
    # ====================================================
    p_tabela2 = doc.add_paragraph()
    run_tab2 = p_tabela2.add_run("Antena GEO")
    run_tab2.bold = True
    p_tabela2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tabela2.paragraph_format.space_after = Pt(6)

    tabela2 = doc.add_table(rows=len(dados_tabela_equip) + 2, cols=5)
    tabela2.alignment = WD_TABLE_ALIGNMENT.CENTER
    adicionar_bordas_tabela(tabela2)

    # Linha 0 (Título Mesclado)
    titulo_cells2 = tabela2.rows[0].cells
    celula_titulo2 = titulo_cells2[0]
    for i in range(1, 5): celula_titulo2.merge(titulo_cells2[i])
    celula_titulo2.text = "ANTENAS VSAT (GEO)"
    celula_titulo2.paragraphs[0].runs[0].bold = True

    # Linha 1 (Cabeçalhos)
    cabecalhos2 = ["EQUIPAMENTO", "FABRICANTE", "MODELO", "S/N°", "IMEI"]
    for i, texto in enumerate(cabecalhos2):
        cell = tabela2.rows[1].cells[i]
        cell.text = texto
        cell.paragraphs[0].runs[0].bold = True

    # Preenchendo os Dados (Antena BR / Antena CBO)
    for i, linha_dados in enumerate(dados_tabela_equip, start=2):
        for col, valor in enumerate(linha_dados):
            tabela2.rows[i].cells[col].text = str(valor)

    # Centralizando TUDO na Tabela 2
    for row in tabela2.rows:
        for cell in row.cells: cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    doc = anexar_fotos_do_bloco(doc, lista_fotos)

    print("✅ Bloco do Sistema GEO gerado com sucesso!")
    return doc