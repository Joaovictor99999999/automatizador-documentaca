from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from automacoes.asbuilt_blocos.config_layout import Layout
from automacoes.asbuilt_blocos.motor_fotos_locais import anexar_fotos_do_bloco

def adicionar_bordas_tabela(tabela):
    """Desenha bordas pretas simples contornando e dividindo a tabela via XML."""
    tblPr = tabela._tbl.tblPr
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)
        
    for borda in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{borda}')
        b.set(qn('w:val'), 'single') 
        b.set(qn('w:sz'), '4')       
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000') 
        tblBorders.append(b)

def gerar_bloco_lte(doc, lista_fotos):
    print("\n==================================================")
    print("📶 BLOCO 8 - SISTEMA LTE (4G/5G)")
    print("==================================================")

    # 1. Quebra de página isolando o capítulo
    doc.add_page_break()
    p_invisivel = doc.add_paragraph()
    p_invisivel.paragraph_format.space_after = Layout.RESPIRO_TOPO_PAGINA

    # 2. Títulos
    p_titulo = doc.add_paragraph()
    run_titulo = p_titulo.add_run("8. Sistema LTE")
    run_titulo.bold = True
    run_titulo.italic = True
    run_titulo.font.size = Layout.TAMANHO_TITULO_PRINCIPAL
    p_titulo.paragraph_format.left_indent = Layout.RECUO_ESQUERDA_NORMAL
    p_titulo.paragraph_format.space_after = Pt(12)

    p_subtitulo = doc.add_paragraph()
    run_sub = p_subtitulo.add_run("Instalação")
    run_sub.bold = True
    p_subtitulo.paragraph_format.left_indent = Layout.RECUO_ESQUERDA_NORMAL
    p_subtitulo.paragraph_format.space_after = Pt(6)

    # 3. Textos Genéricos
    textos = [
        "O sistema LTE possui sua antena instalada na parte externa, posicionada em uma área "
        "superior estratégica da embarcação. A antena externa tem como função a captação do "
        "sinal oriundo da operadora de telefonia móvel, bem como a sua retransmissão.",
        
        "O rádio (modem LTE) foi instalado no rack de telecomunicações e, para o seu funcionamento, "
        "conta com o SIM card (chip) habilitado e fornecido pela contratante. Abaixo, apresenta-se a tabela "
        "com a relação dos equipamentos, materiais e seus respectivos locais de instalação."
    ]

    for idx, frase in enumerate(textos):
        p_texto = doc.add_paragraph(frase)
        p_texto.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_texto.paragraph_format.left_indent = Layout.RECUO_ESQUERDA_NORMAL
        p_texto.paragraph_format.right_indent = Layout.RECUO_DIREITA_NORMAL
        p_texto.paragraph_format.space_after = Pt(18) if idx == len(textos) - 1 else Pt(0)

    # 4. Dados da Tabela LTE
    # Você pode alterar esses dados se algum navio tiver distâncias diferentes
    dados_lte = [
        ("Rack BR", "Área Externa", "Antena Omini", "LMR-400", 1, 40, 40),
        ("Rack BR", "Passadiço", "Antena Omini", "LMR-400", 1, 15, 15),
        ("Rack BR", "Rack BR", "Fortinet FortiGate 60F", "LTE 3G/4G", 0, 0, 0)
    ]

    # 5. Desenhando a Tabela
    tabela = doc.add_table(rows=len(dados_lte) + 2, cols=7)
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    adicionar_bordas_tabela(tabela)

    # Linha 0 (Título Mesclado)
    titulo_cells = tabela.rows[0].cells
    celula_titulo = titulo_cells[0]
    for i in range(1, 7): celula_titulo.merge(titulo_cells[i])
    celula_titulo.text = "INSTALAÇÃO DO SISTEMA LTE - CELULAR 4G/5G"
    celula_titulo.paragraphs[0].runs[0].bold = True

    # Linha 1 (Cabeçalhos)
    cabecalhos = ["PONTO A", "PONTO B", "EQUIPAMENTO", "ESPECIFICAÇÃO", "QTD DE CABOS", "DISTÂNCIA", "TOTAL"]
    for i, texto in enumerate(cabecalhos):
        cell = tabela.rows[1].cells[i]
        cell.text = texto
        cell.paragraphs[0].runs[0].bold = True

    # Injetando Dados
    for i, linha_dados in enumerate(dados_lte, start=2):
        for col, valor in enumerate(linha_dados):
            tabela.rows[i].cells[col].text = str(valor)

    # Centralizando tudo
    for row in tabela.rows:
        for cell in row.cells: cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    doc = anexar_fotos_do_bloco(doc, lista_fotos)

    print("✅ Bloco do Sistema LTE gerado com sucesso!")
    return doc