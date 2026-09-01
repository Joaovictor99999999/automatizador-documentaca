import os
from io import BytesIO
from PIL import Image
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from automacoes.asbuilt_blocos.config_layout import Layout
from automacoes.asbuilt_blocos.motor_fotos_locais import anexar_fotos_do_bloco

def inserir_diagrama_padrao(doc, nome_arquivo, legenda):
    """Insere imagens da pasta de templates, ideal para tabelas gigantes em formato de print."""
    caminho = os.path.join("templates", "imagens_fixas", nome_arquivo)
    if not os.path.exists(caminho):
        # Não achou a imagem? Sem problema, ele apenas avisa e segue em frente (para você colar na mão depois se preferir)
        print(f"⚠️ Imagem '{nome_arquivo}' não encontrada. Espaço deixado para inserção manual.")
        return

    try:
        img = Image.open(caminho)
        img.thumbnail((1920, 1080)) 
        imagem_memoria = BytesIO()
        img.save(imagem_memoria, format="PNG")
        imagem_memoria.seek(0)
        
        p_legenda = doc.add_paragraph()
        p_legenda.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_legenda = p_legenda.add_run(legenda)
        run_legenda.bold = True
        p_legenda.paragraph_format.keep_with_next = True
        p_legenda.paragraph_format.space_after = Pt(6)
        
        p_foto = doc.add_paragraph()
        p_foto.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Largura generosa para tabelas compridas caberem bem na folha
        p_foto.add_run().add_picture(imagem_memoria, width=Inches(6.5))
        doc.add_paragraph("")
        
    except Exception as e:
        print(f"❌ Erro ao processar a imagem {nome_arquivo}: {e}")


def gerar_bloco_cftv(doc, qtd_encodes, qtd_cameras, lista_fotos):
    print("\n==================================================")
    print("📹 BLOCO 7 - SISTEMA CFTV")
    print("==================================================")

    # 1. Coleta das variáveis numéricas

    # 2. Estrutura da Página
    doc.add_page_break()
    p_invisivel = doc.add_paragraph()
    p_invisivel.paragraph_format.space_after = Layout.RESPIRO_TOPO_PAGINA

    p_titulo = doc.add_paragraph()
    run_titulo = p_titulo.add_run("7. Sistema CFTV")
    run_titulo.bold = True
    run_titulo.italic = True
    run_titulo.font.size = Layout.TAMANHO_TITULO_PRINCIPAL
    p_titulo.paragraph_format.left_indent = Layout.RECUO_ESQUERDA_NORMAL
    p_titulo.paragraph_format.space_after = Pt(12)

    p_subtitulo = doc.add_paragraph()
    run_sub = p_subtitulo.add_run("Instalação")
    run_sub.bold = True
    p_subtitulo.paragraph_format.left_indent = Layout.RECUO_ESQUERDA_NORMAL
    p_subtitulo.paragraph_format.space_after = Pt(6)

    # 3. Texto Genérico com as Variáveis Injetadas
    texto_instalacao = (
        f"Da parte do sistema CFTV, foi realizada a instalação de {qtd_cameras} câmeras e {qtd_encodes} "
        "encoders conforme tabela a seguir, em áreas definidas e apontadas in loco. "
        "A identificação seguiu o projeto de rede definido pela Petrobras conforme tabela "
        "apresentada na página 5, foi utilizado o cabo de blindagem CAT 6 F/UTP, para a "
        "instalação de todas as câmeras até o RACK."
    )

    p_texto = doc.add_paragraph(texto_instalacao)
    p_texto.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_texto.paragraph_format.left_indent = Layout.RECUO_ESQUERDA_NORMAL
    p_texto.paragraph_format.right_indent = Layout.RECUO_DIREITA_NORMAL
    p_texto.paragraph_format.space_after = Pt(18)

    # 4. Chamada Opcional das Imagens da Tabela
    # Se você colocar os prints lá na pasta, ele monta sozinho. Se não colocar, ele ignora.
    inserir_diagrama_padrao(doc, "tabela_cftv_1.png", "Tabela de Identificação CFTV - Parte 1")
    inserir_diagrama_padrao(doc, "tabela_cftv_2.png", "Tabela de Identificação CFTV - Parte 2")
    
    doc = anexar_fotos_do_bloco(doc, lista_fotos)

    print("✅ Bloco do Sistema CFTV gerado com sucesso!")
    return doc