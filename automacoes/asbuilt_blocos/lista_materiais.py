import os
from io import BytesIO
from PIL import Image
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from automacoes.asbuilt_blocos.config_layout import Layout
from automacoes.asbuilt_blocos.motor_fotos_locais import anexar_fotos_do_bloco

def inserir_imagem_lista(doc, nome_arquivo):
    caminho = os.path.join("templates", "imagens_fixas", nome_arquivo)
    if not os.path.exists(caminho):
        print(f"⚠️ Imagem '{nome_arquivo}' não encontrada. Espaço deixado no Tópico 11 para inserção manual.")
        return
    try:
        img = Image.open(caminho)
        img.thumbnail((1920, 1080)) 
        imagem_memoria = BytesIO()
        img.save(imagem_memoria, format="PNG")
        imagem_memoria.seek(0)
        
        p_foto = doc.add_paragraph()
        p_foto.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_foto.add_run().add_picture(imagem_memoria, width=Inches(6.5))
        doc.add_paragraph("")
    except Exception as e:
        print(f"❌ Erro ao processar a imagem {nome_arquivo}: {e}")

def gerar_bloco_materiais(doc, lista_fotos):
    print("\n==================================================")
    print("📦 BLOCO 11 - LISTA DE MATERIAIS")
    print("==================================================")

    doc.add_page_break()
    p_invisivel = doc.add_paragraph()
    p_invisivel.paragraph_format.space_after = Layout.RESPIRO_TOPO_PAGINA

    p_titulo = doc.add_paragraph()
    run_titulo = p_titulo.add_run("11. Lista de materiais")
    run_titulo.bold = True
    run_titulo.italic = True
    run_titulo.font.size = Layout.TAMANHO_TITULO_PRINCIPAL
    p_titulo.paragraph_format.left_indent = Layout.RECUO_ESQUERDA_NORMAL
    p_titulo.paragraph_format.space_after = Pt(18)

    doc = anexar_fotos_do_bloco(doc, lista_fotos)
    
    print("✅ Bloco de Lista de Materiais gerado com sucesso!")
    return doc