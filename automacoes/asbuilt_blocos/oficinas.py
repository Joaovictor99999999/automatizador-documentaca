from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from automacoes.asbuilt_blocos.config_layout import Layout

def adicionar_bordas_tabela(tabela):
    tblPr = tabela._tbl.tblPr
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)
    for borda in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{borda}')
        b.set(qn('w:val'), 'single') 
        b.set(qn('w:sz'), '4')       
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000') 
        tblBorders.append(b)

def gerar_bloco_oficinas(doc):
    print("\n==================================================")
    print("🛠️ BLOCO 10 - OFICINAS")
    print("==================================================")

    doc.add_page_break()
    p_invisivel = doc.add_paragraph()
    p_invisivel.paragraph_format.space_after = Layout.RESPIRO_TOPO_PAGINA

    p_titulo = doc.add_paragraph()
    run_titulo = p_titulo.add_run("10. Oficinas")
    run_titulo.bold = True
    run_titulo.italic = True
    run_titulo.font.size = Layout.TAMANHO_TITULO_PRINCIPAL
    p_titulo.paragraph_format.left_indent = Layout.RECUO_ESQUERDA_NORMAL
    p_titulo.paragraph_format.space_after = Pt(12)

    texto_oficinas = (
        "Para realizar as atividades descritas nesse projeto executivo, será necessário as "
        "oficinas de maneira dedicada. A não disponibilização desses recursos a seguir, "
        "comprometerá o cronograma definido."
    )
    p_texto = doc.add_paragraph(texto_oficinas)
    p_texto.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_texto.paragraph_format.left_indent = Layout.RECUO_ESQUERDA_NORMAL
    p_texto.paragraph_format.right_indent = Layout.RECUO_DIREITA_NORMAL
    p_texto.paragraph_format.space_after = Pt(18)

    # MATRIZ DA TABELA DE OFICINAS
    dados_oficinas = [
        ("CALDERARIA", "SOLDA", "Será realizado a fixação do pedestal da caixa metálica no corpo do pedestal, fixação do pedestal da antena Kymeta, criação de pescoço de ganso, criação de bases para fixação das câmeras externas.", "CLIENTE"),
        ("ELÉTRICA", "INSTALAÇÃO", "Após o lançamento do cabo de elétrica (2,5mmx3 vias) a energização será feita diretamente na UPS localizada nos racks.", "OLS"),
        ("MOVIMENTAÇÃO DE CARGAS", "ANTENA E PEDESTAL", "Em função do peso de cada elemento, a movimentação desses materiais deverá ser realizada com apoio de guindaste ou IRATA, além da colocação das câmeras nos pontos que precisam de acesso por cordas.", "CLIENTE"),
        ("ABERTURA DE FORRO", "ABERTURA DE FORRO", "Será necessário uma equipe para realizar a abertura e fechamento do forro para passagem dos cabos de dados e elétrica.", "CLIENTE"),
        ("LANÇAMENTO", "PASSAGEM DE CABO", "Os cabos dos sistemas de GEO, LEO, LTE serão instalados pelo time da OLS. Referente aos outros sistemas, os cabos serão lançados pelo cliente.", "OLS")
    ]

    tabela = doc.add_table(rows=len(dados_oficinas) + 1, cols=4)
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    adicionar_bordas_tabela(tabela)
    
    # 1. Desliga o comportamento automático do Word de tentar adivinhar as larguras
    tabela.autofit = False
    tabela.allow_autofit = False

    cabecalhos = ["DISCIPLINA", "ATIVIDADE", "DESCRITIVO", "AÇÃO"]
    for i, texto in enumerate(cabecalhos):
        cell = tabela.rows[0].cells[i]
        cell.text = texto
        cell.paragraphs[0].runs[0].bold = True

    for i, linha_dados in enumerate(dados_oficinas, start=1):
        for col, valor in enumerate(linha_dados):
            tabela.rows[i].cells[col].text = str(valor)

    # 2. Régua de larguras exatas (Total = ~6.5 polegadas, que é a área útil da folha A4)
    larguras = [Inches(1), Inches(1), Inches(4), Inches(0.8)]

    for row in tabela.rows:
        # Aplica a largura matemática célula por célula para forçar o Word a obedecer
        for i, cell in enumerate(row.cells):
            cell.width = larguras[i]
            
            # Alinhamento do texto
            if i == 2 and row != tabela.rows[0]: # Coluna "Descritivo" alinhada à esquerda
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            else: # Resto centralizado
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    print("✅ Bloco de Oficinas gerado com sucesso!")
    return doc