import os
import copy
from docx import Document
from docx.shared import Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

# A FUNÇÃO AGORA RECEBE OS 'dados' DIRETAMENTE DA INTERFACE
def capturar_e_gerar_cabecalho(dados):
    print("\n==================================================")
    print("📋 PROCESSANDO DADOS DO CABEÇALHO / CAPA")
    print("==================================================")

    caminho_template = os.path.join("templates", "folha1.docx")
    if not os.path.exists(caminho_template):
        print(f"\n❌ Erro: Template '{caminho_template}' não encontrado.")
        return None

    doc = Document(caminho_template)

    try:
        preencher_tabela_cabecalho(doc, dados)
        adicionar_linha_versao(doc, dados)
        escrever_titulo_central(doc, dados)
        print("\n✅ Bloco [Cabeçalho e Capa] preenchido com sucesso!")
    except Exception as e:
        print(f"\n⚠️ Erro ao indexar metadados: {e}")

    # Agora devolvemos apenas o documento, pois a interface já tem os dados
    return doc


def get_unique_cells(row):
    """Retorna as células de uma linha sem repetir células mescladas (horizontal)."""
    seen = set()
    cells = []
    for cell in row.cells:
        key = id(cell._tc)
        if key not in seen:
            seen.add(key)
            cells.append(cell)
    return cells


def set_cell_text(cell, texto):
    """Escreve o texto no primeiro parágrafo da célula, mantendo a formatação
    do primeiro run existente (fonte/negrito) e limpando parágrafos extras."""
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = texto
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.text = texto
    for extra_p in cell.paragraphs[1:]:
        extra_p.text = ""


def preencher_tabela_cabecalho(doc, dados):
    """Preenche a tabela de cabeçalho escrevendo cada valor na célula
    VIZINHA ao label (e não dentro da própria célula do label)."""
    tabela = doc.tables[0]

    for row in tabela.rows:
        cells = get_unique_cells(row)
        for idx, cell in enumerate(cells):
            texto = cell.text.strip()
            if idx + 1 >= len(cells):
                continue
            valor_cell = cells[idx + 1]

            if texto.startswith(("N° PROJETO", "Nº PROJETO")):
                set_cell_text(valor_cell, dados["num_projeto"])
            elif texto.startswith("FOLHA:"):
                set_cell_text(valor_cell, dados["folha"])
            elif texto.startswith("EMITENTE:"):
                set_cell_text(valor_cell, dados["emitente"])
            elif texto.startswith("REV:"):
                set_cell_text(valor_cell, dados["rev"])
            elif texto.startswith(("GERÊNCIA DO CONTRATO", "GERENCIA DO CONTRATO")):
                set_cell_text(valor_cell, dados["gerencia"])
            elif texto.startswith(("TÍTULO:", "TITULO:")):
                set_cell_text(valor_cell, dados["titulo"])
            elif texto == "DATA:":
                set_cell_text(valor_cell, dados["data"])


def adicionar_linha_versao(doc, dados):
    """Duplica a última linha da tabela (TÍTULO/DATA) para criar a linha
    de VERSÃO / DATA DA ÚLTIMA MODIFICAÇÃO, preservando os merges originais."""
    tabela = doc.tables[0]
    tbl = tabela._tbl

    ultima_tr = tabela.rows[-1]._tr
    nova_tr = copy.deepcopy(ultima_tr)
    tbl.append(nova_tr)

    tabela = doc.tables[0]
    nova_row = tabela.rows[-1]
    cells = get_unique_cells(nova_row)
    
    set_cell_text(cells[0], f"VERSÃO {dados['versao']}")
    set_cell_text(cells[1], "")
    set_cell_text(cells[2], "DATA DA\nÚLTIMA\nMODIFICAÇÃO")
    set_cell_text(cells[3], dados["data_modificacao"])


def paragrafos_depois_da_tabela(doc, tabela):
    """Retorna, em ordem, apenas os parágrafos (<w:p>) que vêm DEPOIS da
    tabela informada no corpo do documento."""
    tbl_element = tabela._tbl
    body = doc.element.body
    depois = False
    paragrafos = []
    for child in body:
        if child is tbl_element:
            depois = True
            continue
        if depois and child.tag == qn("w:p"):
            paragrafos.append(Paragraph(child, doc))
    return paragrafos


def escrever_titulo_central(doc, dados):
    """Escreve o bloco de título grande no meio da capa."""
    tabela = doc.tables[0]
    candidatos = paragrafos_depois_da_tabela(doc, tabela)
    paragrafos_vazios = [p for p in candidatos if p.text.strip() == ""]

    linhas = [dados["titulo_meio_1"], dados["titulo_meio_2"], dados["titulo_meio_3"]]

    while len(paragrafos_vazios) < len(linhas):
        paragrafos_vazios.append(doc.add_paragraph())

    for i, (p, texto) in enumerate(zip(paragrafos_vazios, linhas)):
        for r in list(p.runs):
            r.text = ""
        run = p.add_run(texto) if not p.runs else p.runs[0]
        run.text = texto
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.italic = True
        run.font.name = "Arial"
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Emu(541020)
        p.paragraph_format.space_after = Pt(24)
        if i == 0:
            p.paragraph_format.space_before = Pt(120)