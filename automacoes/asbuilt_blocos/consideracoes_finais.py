from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from automacoes.asbuilt_blocos.config_layout import Layout

def gerar_bloco_consideracoes(doc):
    print("\n==================================================")
    print("🏁 BLOCO 12 - CONSIDERAÇÕES FINAIS")
    print("==================================================")

    doc.add_page_break()
    p_invisivel = doc.add_paragraph()
    p_invisivel.paragraph_format.space_after = Layout.RESPIRO_TOPO_PAGINA

    p_titulo = doc.add_paragraph()
    run_titulo = p_titulo.add_run("12. CONSIDERAÇÕES FINAIS")
    run_titulo.bold = True
    run_titulo.italic = True
    run_titulo.font.size = Layout.TAMANHO_TITULO_PRINCIPAL
    p_titulo.paragraph_format.left_indent = Layout.RECUO_ESQUERDA_NORMAL
    p_titulo.paragraph_format.space_after = Pt(12)

    frases = [
        "O documento elaborado tem como objetivo ser utilizado na atividade de adequação ao Anexo III seguindo a TIC_ET-0600.00-5510-760-PPT-542.",
        "Está contido nesse documento todas as disciplinas que devem ser tratadas e implantadas no projeto, além das especificações de cada atividade, localização de cada sistema e materiais que serão utilizados.",
        "Também estão contidos nesse documento, via anexo, os unifilares de rede, além da ART (Anotação de Responsabilidade Técnica) que atende a uma exigência estabelecida pelo CREA que; para atividades técnicas, conforme a Lei nº 6.496/77, o documento deve ser registrado via portal.",
        "O documento em questão pode a qualquer momento passar por retificações em função de ajuste estrutural ou até mesmo por mudança de ET."
    ]

    for idx, frase in enumerate(frases):
        p_texto = doc.add_paragraph(frase)
        p_texto.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_texto.paragraph_format.left_indent = Layout.RECUO_ESQUERDA_NORMAL
        p_texto.paragraph_format.right_indent = Layout.RECUO_DIREITA_NORMAL
        p_texto.paragraph_format.space_after = Pt(18) if idx == len(frases) - 1 else Pt(12)

    print("✅ Bloco de Considerações Finais gerado com sucesso!")
    return doc