from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from automacoes.asbuilt_blocos.config_layout import Layout
from automacoes.asbuilt_blocos.motor_fotos_locais import anexar_fotos_do_bloco

def adicionar_bordas_tabela(tabela):
    """Desenha bordas pretas simples contornando e dividindo a tabela via XML."""
    tblPr = tabela._tbl.tblPr
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)
        
    for borda in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{borda}')
        b.set(qn('w:val'), 'single') 
        b.set(qn('w:sz'), '4')       
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000') 
        tblBorders.append(b)

def gerar_bloco_leo(doc, qtd_antenas, lista_fotos):
    print("\n==================================================")
    print("🛰️ BLOCO 5 - SISTEMA LEO (KYMETA / STARLINK)")
    print("==================================================")
    # ====================================================
    # PÁGINA 1: Título + Tabela 1 + Textos de Instalação
    # ====================================================
    doc.add_page_break()
    p_invisivel = doc.add_paragraph()
    p_invisivel.paragraph_format.space_after = Layout.RESPIRO_TOPO_PAGINA

    p_titulo = doc.add_paragraph()
    run_titulo = p_titulo.add_run("5. Sistema LEO")
    run_titulo.bold = True
    run_titulo.italic = True
    run_titulo.font.size = Layout.TAMANHO_TITULO_PRINCIPAL
    p_titulo.paragraph_format.left_indent = Layout.RECUO_ESQUERDA_NORMAL
    p_titulo.paragraph_format.space_after = Pt(12)

    p_subtitulo = doc.add_paragraph()
    run_sub = p_subtitulo.add_run("Instalação")
    run_sub.bold = True
    p_subtitulo.paragraph_format.left_indent = Layout.RECUO_ESQUERDA_NORMAL
    p_subtitulo.paragraph_format.space_after = Pt(6)

    # 1. PREPARAÇÃO DAS MATRIZES DINÂMICAS (Para as duas tabelas)
    dados_antena_1 = [
        ("Rack CBO", "Tijupa BB", "Antena.Kymeta", "CAT6 Blindado", 2, 20, 40),
        ("Antena", "QDG_BB", "Fonte Kymeta", "Elet.PP 2,5x3mm", 1, 35, 35),
        ("UPS Rack", "Antena BB", "Antena.Kymeta", "Elet.PP 2,5x3mm", 1, 20, 20)
    ]
    
    dados_antena_2 = [
        ("Novo Rack BR", "Tijupa BE", "Antena.Kymeta", "CAT6 Blindado", 2, 35, 70),
        ("Antena", "QDG_BE", "Fonte Kymeta", "Elet.PP 2,5x3mm", 1, 2, 2),
        ("UPS Rack", "Antena BE", "Antena.Kymeta", "Elet.PP 2,5x3mm", 1, 35, 35)
    ]

    if qtd_antenas == "2":
        texto_instalacao = (
            "O sistema LEO teve suas antenas instaladas preferencialmente em áreas diferentes "
            "(em extremidades opostas) a fim de mitigar problemas de áreas de sombra. "
            "Também é de extrema importância ressaltar que a instalação foi realizada "
            "pelo corpo técnico da OLS."
        )
        dados_tabela_cabeamento = dados_antena_1 + dados_antena_2
        
        # Matriz da Tabela 2 (Equipamentos Kymeta)
        dados_tabela_equip = [
            ("Antena A", "KYMETA", "Peregrine U8", "xxxxxxx", "xxxxxxx"),
            ("Antena B", "KYMETA", "Peregrine U8", "xxxxxxx", "xxxxxxx")
        ]
    else:
        texto_instalacao = (
            "O sistema LEO foi instalado utilizando apenas uma antena, posicionada no ponto "
            "mais alto da estrutura. O local da instalação garante total visada da antena, "
            "porém fica ressaltado que a área apresenta dificuldade de acesso para futuras manutenções."
        )
        dados_tabela_cabeamento = [
            ("Rack Telecom", "Topo da Torre", "Antena.Kymeta", "CAT6 Blindado", 2, 20, 40),
            ("Antena", "QDG", "Fonte Kymeta", "Elet.PP 2,5x3mm", 1, 35, 35),
            ("UPS Rack", "Antena LEO", "Antena.Kymeta", "Elet.PP 2,5x3mm", 1, 20, 20)
        ]
        
        # Matriz da Tabela 2 (Equipamentos Kymeta)
        dados_tabela_equip = [
            ("Antena", "KYMETA", "Peregrine U8", "xxxxxxx", "xxxxxxx")
        ]

    # TEXTO DE INSTALAÇÃO
    p_texto = doc.add_paragraph(texto_instalacao)
    p_texto.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_texto.paragraph_format.left_indent = Layout.RECUO_ESQUERDA_NORMAL
    p_texto.paragraph_format.right_indent = Layout.RECUO_DIREITA_NORMAL
    p_texto.paragraph_format.space_after = Pt(18)

    # TABELA 1: CABEAMENTO
    tabela1 = doc.add_table(rows=len(dados_tabela_cabeamento) + 3, cols=7)
    tabela1.alignment = WD_TABLE_ALIGNMENT.CENTER
    adicionar_bordas_tabela(tabela1)

    titulo_cells = tabela1.rows[0].cells
    celula_titulo = titulo_cells[0]
    for i in range(1, 7): celula_titulo.merge(titulo_cells[i])
    celula_titulo.text = "INSTALAÇÃO DO SISTEMA LEO"
    celula_titulo.paragraphs[0].runs[0].bold = True

    cabecalhos1 = ["PONTO A", "PONTO B", "EQUIPAMENTO", "ESPECIFICAÇÃO", "QTD DE CABOS", "DISTÂNCIA", "TOTAL"]
    for i, texto in enumerate(cabecalhos1):
        cell = tabela1.rows[1].cells[i]
        cell.text = texto
        cell.paragraphs[0].runs[0].bold = True

    for i, linha_dados in enumerate(dados_tabela_cabeamento, start=2):
        for col, valor in enumerate(linha_dados):
            tabela1.rows[i].cells[col].text = str(valor)

    soma_qtd = sum(item[4] for item in dados_tabela_cabeamento)
    soma_dist = sum(item[5] for item in dados_tabela_cabeamento)
    soma_total = sum(item[6] for item in dados_tabela_cabeamento)

    linha_rodape = tabela1.rows[-1].cells
    celula_rodape_texto = linha_rodape[0]
    for i in range(1, 4): celula_rodape_texto.merge(linha_rodape[i])
    celula_rodape_texto.text = "TOTAL GERAL (m)"
    celula_rodape_texto.paragraphs[0].runs[0].bold = True
    
    linha_rodape[4].text = str(soma_qtd)
    linha_rodape[4].paragraphs[0].runs[0].bold = True
    linha_rodape[5].text = str(soma_dist)
    linha_rodape[5].paragraphs[0].runs[0].bold = True
    linha_rodape[6].text = str(soma_total)
    linha_rodape[6].paragraphs[0].runs[0].bold = True

    for row in tabela1.rows:
        for cell in row.cells: cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("") # Respiro após a tabela 1

    # TEXTOS EXTRAS (AINDA NA PÁGINA 1)
    texto_especificacoes = (
        "No sistema VSAT LEO, a habilitação e a ativação da antena junto à OneWeb foi "
        "realizada, e a confecção e instalação do pedestal seguem rigorosamente o "
        "projeto enviado. Para a antena Kymeta, todos os elementos de conexão ficam "
        "localizados dentro da própria estrutura. A interface de conexão ocorre por "
        "meio de dois cabos blindados (SFTP) — um de dados e outro de gerência —, "
        "que chegam até o rack e estão devidamente identificados em ambas as pontas. "
        "O sistema de alimentação da antena divide-se entre as partes externa e interna:"
    )
    p_esp = doc.add_paragraph(texto_especificacoes)
    p_esp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_esp.paragraph_format.left_indent = Layout.RECUO_ESQUERDA_NORMAL
    p_esp.paragraph_format.right_indent = Layout.RECUO_DIREITA_NORMAL

    p_externa = doc.add_paragraph()
    p_externa.paragraph_format.left_indent = Layout.RECUO_ESQUERDA_TOPICO
    p_externa.paragraph_format.right_indent = Layout.RECUO_DIREITA_NORMAL
    p_externa.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_externa.add_run("•  Externa: ").bold = True
    p_externa.add_run("A antena possui uma fonte de conversão de energia AC/DC acomodada em uma "
                      "caixa metálica (QDG). Esta caixa é dotada de um disjuntor bipolar de 10A, adaptadores "
                      "de rede, identificações, fixação e organização de acordo com a figura 01.")

    p_interna = doc.add_paragraph()
    p_interna.paragraph_format.left_indent = Layout.RECUO_ESQUERDA_TOPICO
    p_interna.paragraph_format.right_indent = Layout.RECUO_DIREITA_NORMAL
    p_interna.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_interna.paragraph_format.space_after = Pt(18)
    p_interna.add_run("•  Interna: ").bold = True
    p_interna.add_run("Os cabos de dados (data e Adm) partem diretamente da antena e chegam até "
                      "o rack de Telecom. Eles estão devidamente identificados em ambas as pontas e "
                      "suas conexões seguem as diretrizes do diagrama unifilar (conforme o anexo-II).")


    # ====================================================
    # PÁGINA 2: Sombra, Orientação e Tabela de Equipamentos
    # ====================================================
    doc.add_page_break()
    
    p_invisivel2 = doc.add_paragraph()
    p_invisivel2.paragraph_format.space_after = Layout.RESPIRO_TOPO_PAGINA

    # Subtítulo: Área de Sombra
    p_sub_sombra = doc.add_paragraph()
    run_sub_sombra = p_sub_sombra.add_run("Área de sombra")
    run_sub_sombra.bold = True
    p_sub_sombra.paragraph_format.left_indent = Layout.RECUO_ESQUERDA_NORMAL
    p_sub_sombra.paragraph_format.space_after = Pt(6)

    texto_sombra = (
        "A instalação da antena precisa seguir as orientações informadas pelo fabricante conforme "
        "manual no que diz respeito à proximidade a outros equipamentos, principalmente ao radar, "
        "conforme figura a seguir."
    )
    p_sombra = doc.add_paragraph(texto_sombra)
    p_sombra.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_sombra.paragraph_format.left_indent = Layout.RECUO_ESQUERDA_NORMAL
    p_sombra.paragraph_format.right_indent = Layout.RECUO_DIREITA_NORMAL
    p_sombra.paragraph_format.space_after = Pt(18)

    # Subtítulo: Orientação do Fabricante
    p_sub_ori = doc.add_paragraph()
    run_sub_ori = p_sub_ori.add_run("Orientação do fabricante:")
    run_sub_ori.bold = True
    p_sub_ori.paragraph_format.left_indent = Layout.RECUO_ESQUERDA_NORMAL
    p_sub_ori.paragraph_format.space_after = Pt(6)

    texto_ori = (
        "O UT deve ser instalado com um nível de ±2°. Certifique-se de que o UT tenha uma linha "
        "de visão clara para o céu. Linha de visão clara para o céu significa que o UT não tem nenhuma "
        "obstrução de lado a lado a 37° acima do horizonte em um círculo de 360°. Se uma linha de visão "
        "clara não estiver disponível, mude para um local onde a linha de visão seja clara. "
        "Se não for possível identificar um local de instalação adequado de acordo com estas diretrizes, "
        "fazer o melhor esforço para acomodar ajudará a minimizar as interrupções do serviço."
    )
    p_ori = doc.add_paragraph(texto_ori)
    p_ori.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_ori.paragraph_format.left_indent = Layout.RECUO_ESQUERDA_NORMAL
    p_ori.paragraph_format.right_indent = Layout.RECUO_DIREITA_NORMAL
    p_ori.paragraph_format.space_after = Pt(18)

    # TABELA 2: ANTENA KYMETA (Equipamentos)
    p_tabela2 = doc.add_paragraph()
    run_tab2 = p_tabela2.add_run("Antena Kymeta")
    run_tab2.bold = True
    p_tabela2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tabela2.paragraph_format.space_after = Pt(6)

    tabela2 = doc.add_table(rows=len(dados_tabela_equip) + 2, cols=5)
    tabela2.alignment = WD_TABLE_ALIGNMENT.CENTER
    adicionar_bordas_tabela(tabela2)

    # Linha 0 (Título Mesclado)
    titulo_cells2 = tabela2.rows[0].cells
    celula_titulo2 = titulo_cells2[0]
    for i in range(1, 5): celula_titulo2.merge(titulo_cells2[i])
    celula_titulo2.text = "ANTENA SDWAN (ONEWB)"
    celula_titulo2.paragraphs[0].runs[0].bold = True

    # Linha 1 (Cabeçalhos)
    cabecalhos2 = ["EQUIPAMENTO", "FABRICANTE", "MODELO", "S/N°", "IMEI"]
    for i, texto in enumerate(cabecalhos2):
        cell = tabela2.rows[1].cells[i]
        cell.text = texto
        cell.paragraphs[0].runs[0].bold = True

    # Preenchendo os Dados Genéricos "xxxxxxx"
    for i, linha_dados in enumerate(dados_tabela_equip, start=2):
        for col, valor in enumerate(linha_dados):
            tabela2.rows[i].cells[col].text = str(valor)

    # Centralizando TUDO na Tabela 2
    for row in tabela2.rows:
        for cell in row.cells: cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc = anexar_fotos_do_bloco(doc, lista_fotos)

    print("✅ Bloco do Sistema LEO (com as duas Tabelas Dinâmicas) gerado com sucesso!")
    return doc