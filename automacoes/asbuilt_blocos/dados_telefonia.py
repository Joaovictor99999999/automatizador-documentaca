import os
from io import BytesIO
from PIL import Image
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from automacoes.asbuilt_blocos.config_layout import Layout
from automacoes.asbuilt_blocos.motor_fotos_locais import anexar_fotos_do_bloco

def adicionar_bordas_tabela(tabela):
    """Desenha bordas pretas simples contornando e dividindo a tabela via XML, blindando contra templates sem estilos."""
    tblPr = tabela._tbl.tblPr
    # Verifica se já existe uma configuração de borda, se não, cria
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)
        
    for borda in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{borda}')
        b.set(qn('w:val'), 'single') # Linha simples
        b.set(qn('w:sz'), '4')       # Espessura 1/2 pt
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000') # Cor preta
        tblBorders.append(b)

def inserir_diagrama_padrao(doc, nome_arquivo, legenda):
    """
    Insere diagrama estático com legenda superior em negrito e centralizada nativamente.
    Sempre começa em página nova, com o mesmo 'respiro' usado nos títulos (RESPIRO_TOPO_PAGINA),
    para a legenda/imagem não colidir com a logo do cabeçalho.
    """
    caminho = os.path.join("templates", "imagens_fixas", nome_arquivo)

    if not os.path.exists(caminho):
        print(f"⚠️ Aviso: Imagem padrão '{nome_arquivo}' não encontrada em templates/imagens_fixas/")
        return

    try:
        img = Image.open(caminho)
        img.thumbnail((1920, 1080))

        imagem_memoria = BytesIO()
        img.save(imagem_memoria, format="PNG")
        imagem_memoria.seek(0)

        # === MESMO HACK DO adicionar_titulo: força página nova + respiro do topo ===
        doc.add_page_break()
        p_invisivel = doc.add_paragraph()
        p_invisivel.paragraph_format.space_after = Layout.RESPIRO_TOPO_PAGINA

        # 1. PARÁGRAFO DA LEGENDA (em cima, centralizada e em negrito)
        p_legenda = doc.add_paragraph()
        p_legenda.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_legenda = p_legenda.add_run(legenda)
        run_legenda.bold = True

        # Legenda nunca desgruda da foto se houver quebra de página
        p_legenda.paragraph_format.keep_with_next = True
        p_legenda.paragraph_format.space_after = Pt(7)

        # 2. PARÁGRAFO DA FOTO (centralizado nativamente na folha, sem tabela)
        p_foto = doc.add_paragraph()
        p_foto.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_foto.add_run().add_picture(imagem_memoria, width=Inches(6.5))

        # Respiro longo após a imagem
        doc.add_paragraph("")

    except Exception as e:
        print(f"❌ Erro ao processar o diagrama {nome_arquivo}: {e}")

def gerar_bloco_telefonia(doc, lista_fotos):
    print("\n==================================================")
    print("📞 BLOCO 4 - SISTEMA DE DADOS E TELEFONIA")
    print("==================================================")

    doc.add_page_break()

    p_invisivel = doc.add_paragraph()
    p_invisivel.paragraph_format.space_after = Layout.RESPIRO_TOPO_PAGINA

    p_titulo = doc.add_paragraph()
    run_titulo = p_titulo.add_run("4. Sistema de dados e telefonia")
    run_titulo.bold = True
    run_titulo.italic = True
    run_titulo.font.size = Layout.TAMANHO_TITULO_PRINCIPAL
    p_titulo.paragraph_format.left_indent = Layout.RECUO_ESQUERDA_NORMAL
    p_titulo.paragraph_format.space_after = Pt(12)

    p_chamada = doc.add_paragraph()
    run_chamada = p_chamada.add_run("Tabela de equipamentos do Rack de Telecom e Telefonia:")
    run_chamada.bold = True
    p_chamada.paragraph_format.left_indent = Layout.RECUO_ESQUERDA_NORMAL
    p_chamada.paragraph_format.space_after = Pt(6)

    # === TABELA PADRÃO DO RACK: título mesclado + cabeçalho + dados ===
    # Ajuste esses dados por projeto se precisar; a estrutura (4 colunas + título mesclado) é a padrão.
    dados_rack = [
        ("ROUTER", "JUNIPER", "SRX300", "H.A"),
        ("SWITCHER", "CISCO", "2960 24P-A", "H.A"),
        ("PROBE", "PROBE", "PROBE", "H.A"),
        ("SWITCHER", "CISCO", "9300 24P-A", "BR"),
        ("FIRWALL", "FORTGATE", "F-80F", "BR"),
        ("CABO UTP", "CAT6", "LSZH", "REDE"),
        ("CABO SFTP", "CAT6", "BLINDADO", "REDE"),
        ("PATCH CORD", "CAT6", "UTP", "RACK"),
        ("PATCH CORD", "CAT6", "LSZH", "RACK"),
        ("CÂMERA", "INTELBRAS", "VIP 5232 SD IA", "BARCO"),
        ("UPS", "APC", "SENOIDAL", "RACK"),
    ]

    tabela = doc.add_table(rows=2 + len(dados_rack), cols=4)
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    adicionar_bordas_tabela(tabela)

    # Linha 0: título mesclado nas 4 colunas
    titulo_cells = tabela.rows[0].cells
    celula_titulo = titulo_cells[0].merge(titulo_cells[1]).merge(titulo_cells[2]).merge(titulo_cells[3])
    celula_titulo.text = "RACK TELECOM_SISTEMA_ H.A"
    celula_titulo.paragraphs[0].runs[0].bold = True

    # Linha 1: cabeçalho das colunas
    hdr_cells = tabela.rows[1].cells
    for cell, texto in zip(hdr_cells, ["EQUIP", "FABRICANTE", "MODELO", "APLICAÇÃO"]):
        cell.text = texto
        cell.paragraphs[0].runs[0].bold = True

    # Linhas de dados
    for i, (equip, fab, modelo, aplic) in enumerate(dados_rack, start=2):
        linha = tabela.rows[i].cells
        linha[0].text = equip
        linha[1].text = fab
        linha[2].text = modelo
        linha[3].text = aplic

    for row in tabela.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    # 5. Texto de Instalação Padrão
    p_subtitulo = doc.add_paragraph()
    run_sub = p_subtitulo.add_run("Instalação")
    run_sub.bold = True
    p_subtitulo.paragraph_format.left_indent = Layout.RECUO_ESQUERDA_NORMAL
    p_subtitulo.paragraph_format.space_after = Pt(6)

    # Cada frase agora é um parágrafo próprio (sem \n), com space_after = 0 entre elas
    # para ficarem "embaixo uma da outra" sem pular linha.
    frases_instalacao = [
        "Para a rede de dados, foi realizada a instalação de todos os pontos de rede. "
        "A embarcação possui um total de 40 pontos mapeados em sua área interna, "
        "utilizando cabeamento UTP CAT 6 LSZH, distribuídos conforme a tabela a seguir.",
        "Quanto à identificação, o uso das tags enviadas juntamente com os diagramas foi "
        "seguido de maneira obrigatória.",
        "Para o novo sistema de H.A., foi instalado um rack adicional que acomoda todos "
        "os equipamentos H.A. OLS, sendo interligado com o rack BR.",
    ]

    for idx, frase in enumerate(frases_instalacao):
        p_texto = doc.add_paragraph(frase)
        p_texto.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_texto.paragraph_format.left_indent = Layout.RECUO_ESQUERDA_NORMAL
        p_texto.paragraph_format.right_indent = Layout.RECUO_DIREITA_NORMAL
        ultima = idx == len(frases_instalacao) - 1
        p_texto.paragraph_format.space_after = Pt(18) if ultima else Pt(0)

    inserir_diagrama_padrao(doc, "esquematico.png", "Tabela do Esquemático de Conexões")
    inserir_diagrama_padrao(doc, "topologia.png", "Topologia Lógica da Rede")
    
    doc = anexar_fotos_do_bloco(doc, lista_fotos)

    print("✅ Bloco de Telefonia e Dados gerado com sucesso!")
    return doc