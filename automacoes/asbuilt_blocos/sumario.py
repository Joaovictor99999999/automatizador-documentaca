from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def gerar_sumario_dinamico(doc, escopo):
    print("📝 Montando Sumário dinâmico...")

    # 1. FORÇA A QUEBRA DE PÁGINA ANTES DE COMEÇAR (Isola a Capa)
    doc.add_page_break()

    # 2. PULA LINHAS NO TOPO DA FOLHA 2 (Respiro superior)
    doc.add_paragraph("")
    doc.add_paragraph("")

    # 3. TÍTULO MAIOR E CENTRALIZADO
    p_titulo = doc.add_paragraph("SUMÁRIO")
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titulo = p_titulo.runs[0]
    run_titulo.bold = True
    run_titulo.font.size = Pt(22)  # Aumentado para dar destaque
    
    # Mais respiro depois do título
    doc.add_paragraph("")
    doc.add_paragraph("")

    contador = 1

    def adicionar_item(numero, texto):
        p = doc.add_paragraph()
        
        # Empurra o bloco de itens para o centro da folha (Recuo esquerdo)
        p.paragraph_format.left_indent = Pt(50) 
        p.paragraph_format.space_after = Pt(12) # Espaçamento entre as linhas
        
        run_texto = p.add_run(f"{numero}. {texto} ")
        run_texto.bold = True
        
        qtd_pontos = 120 - len(texto) - len(str(numero))
        if qtd_pontos < 5: qtd_pontos = 5
        
        p.add_run("." * qtd_pontos)
        p.add_run(" X").bold = True

    # --- ITENS FIXOS DE INTRODUÇÃO ---
    adicionar_item(contador, "Descrição / Escopo")
    contador += 1
    adicionar_item(contador, "Plantas e desenhos técnicos")
    contador += 1
    adicionar_item(contador, "Especificações técnicas")
    contador += 1

    # --- ITENS DINÂMICOS ---
    if escopo["telefonia"]:
        adicionar_item(contador, "Sistema de dados e telefonia")
        contador += 1
    if escopo["leo"]:
        adicionar_item(contador, "Sistema LEO (Kymeta/Starlink)")
        contador += 1
    if escopo["geo"]:
        adicionar_item(contador, "Sistema GEO (Intellian/Sailor)")
        contador += 1
    if escopo["cftv"]:
        adicionar_item(contador, "Sistema CFTV")
        contador += 1
    if escopo["lte"]:
        adicionar_item(contador, "Sistema LTE")
        contador += 1

    # --- ITENS FIXOS FINAIS ---
    adicionar_item(contador, "Cronograma de Execução")
    contador += 1
    adicionar_item(contador, "Oficinas")
    contador += 1
    adicionar_item(contador, "Lista de materiais")
    contador += 1
    adicionar_item(contador, "Considerações finais")

    # Força a quebra de página para o próximo bloco (Folha 3)
    doc.add_page_break() 
    
    print("✅ Bloco [Sumário] gerado com sucesso!")
    return doc